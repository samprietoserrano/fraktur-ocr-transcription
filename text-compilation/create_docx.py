from datetime import datetime
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
import os
import subprocess

from docx_creation_helpers import add_hyperlink, add_image_to_page, apply_formatting, extract_number, \
    find_matching_path, get_next_filename, include_image_files, set_number_of_columns

def load_image_pages():
    """Returns the page number and path of pages identified as holding an image/visual."""

    image_folder = '/Users/USER/Documents/CESTA-Summer/all-pages-as-jpeg'

    img_nums = [1, 2, 5, 8, 17, 76, 140, 170, 177, 192, 201, 210, 218, 236, 240, 
                        456, 473, 491, 519, 523, 529, 557, 569, 576, 590, 600, 
                        626, 647, 718]
    tbl_nums = [404, 405, 406, 407, 408]

    page_nums = sorted(img_nums+tbl_nums)
    
    paths = dict()
    for number in page_nums:
        paths[number] = find_matching_path(image_folder, number-1) # minus 1 because jpg folder is 0-indexed

    return page_nums, paths, tbl_nums


def load_two_column_pages():
    """Returns the pages identified as having two-columns of text."""

    single_column_pages = [i for i in range(0, 23)]
    single_column_pages += [31, 54, 55, 727, 728, 734]

    pages = []
    for i in range(1, 1000):
        if i not in single_column_pages:
            pages.append(i)
    return pages


def add_intro_page(doc, img_available):
    """Adds an introduction page to the document.
    Requires to be edited to change the content as well as the links and image."""

    # Title heading
    title = doc.add_heading('Caput Bonae Spei hodiernum, das ist: \
                        vollständige Beschreibung des africanischen Vorgebürges der Guten Hofnung', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Subheading
    subheading = doc.add_heading('by Peter Kolb, machine-readable version by Samuel Prieto Serrano', level=1)
    subheading.paragraph_format.space_before = Pt(6)
    subheading.paragraph_format.space_after = Pt(12)

    # Introductory paragraph
    intro_par = doc.add_paragraph(
        "This Docx file is a non-final machine-readable transcription of Peter Kolb's 1719 publication. This "
        "project has been commandeered by the Early Cape Travelers research project at Stanford University. "
        "The file contains every non-blank page of Kolb's book and was created with minimal formatting. Words "
        "with still-unknown spelling or meaning are colored red."
    )
    intro_par.paragraph_format.space_after = Pt(12)
    intro_par.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Second paragraph
    second_par = doc.add_paragraph(
        "Below you can find links to the main resources of this project.\nDue to possible compatibility issues using "
        "hyperlinks, the full link is also included."
    )
    second_par.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    links = [
        ('CESTA Background', 'https://cesta.stanford.edu/research/early-cape-travelers'),
        ('CESTA Project Article', 'https://cesta-io.stanford.edu/anthology/2024-research-anthology/early-cape-travelers/'),
        ('GitHub Repository for Project', 'https://github.com/cesta-online/prj-early-cape-travelers'),
        ('Research Contact: Grant Parker', 'mailto:grparker@stanford.edu')
    ]
    item= 1
    second_par.add_run("\n\n")  # New line    
    for text, url in links:
        second_par.add_run(str(item) + ') ')  # New line
        add_hyperlink(doc, second_par, text, url)
        # second_par.add_run(', (' + url + ')\n')  # New line
        second_par.add_run(', ' + url + '\n')  # New line
        item += 1

    if img_available:
        # Image paragraph
        image_path = '/Users/USER/Documents/CESTA-Summer/all-pages-as-jpeg/images-only/bub_gb_MbxYAAAAcAAJ_0016.jpg'
        image_par = doc.add_paragraph()
        run = image_par.add_run()
        run.add_picture(image_path, width=Inches(1.5), height=Inches(2.5))
        image_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Caption paragraph
        caption_par = doc.add_paragraph()
        caption_run = caption_par.add_run('Portrait of Peter Kolb (1727). Page 17 in the book.')
        caption_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        caption_run.italic = True 
        caption_par.paragraph_format.space_after = Pt(12)


def txt_to_word(folder, output_folder, format, parts=False, img_available=True):
    """Creates either a simple or formatted docx file with the provided txt files."""

    doc_pages = 0

    # Make folder for partial documents
    if parts:
        output_folder = os.path.join(output_folder, "partial_docx")
        os.makedirs(output_folder, exist_ok=True)

    # Create a new Document and get list of txt files
    doc = Document()
    add_intro_page(doc, img_available)
    doc_pages += 1
    
    page_files = [f for f in os.listdir(folder) if f.endswith(".txt")]

    if format:
        two_column_pages = load_two_column_pages()
        if img_available:
            img_tbl_nums, img_paths, tbl_nums = load_image_pages()
            page_files = include_image_files(page_files, img_tbl_nums)

    # Sort the files numerically based on the extracted number
    page_files.sort(key=extract_number)

    pages_added = [] # used for computations only, not saved as output
    for i in range(len(page_files)):
        filename = page_files[i]
        curr_file_number = extract_number(filename)

        # if len(pages_added) > 80:
        #     break

        if parts:
            if doc_pages % 100 == 0:
                # Save the partial document
                today = datetime.today()
                date_string = today.strftime('%Y-%m-%d')

                base_name = 'format-compiled-doc-' if format else 'simple-compiled-doc-'
                next_file = get_next_filename(output_folder, base_name + date_string, 'docx')
                output_file = os.path.join(output_folder, next_file)
                doc.save(output_file)

                print("The partial Docx was saved!")

                # Create a new Document
                doc = Document()
                add_intro_page(doc, img_available)
                doc_pages += 1
        
        if img_available:
            if format and curr_file_number in img_tbl_nums:
                if filename not in pages_added:
                    type = "img" if curr_file_number not in tbl_nums else "tbl"
                    add_image_to_page(doc, img_paths[int(curr_file_number)], curr_file_number, type=type)
                    pages_added.append(filename)
                    doc_pages += 1
                    continue

        with open(os.path.join(folder, filename), 'r', encoding='utf-8') as file:
            content = file.read()
        
        _ = doc.add_section(WD_SECTION.NEW_PAGE)
        doc_pages += 1
        doc.add_heading(filename, level=1)
        if format:
            if curr_file_number in two_column_pages:
                curr_section = doc.add_section(WD_SECTION.CONTINUOUS)
                set_number_of_columns(curr_section, 2)        

            paragraph = doc.add_paragraph()
            apply_formatting(paragraph, content)

            pages_added.append(filename)
            # _ = doc.add_section(WD_SECTION.NEW_PAGE)
        else:
            doc.add_paragraph(content)            
            pages_added.append(filename)


    # Save the document
    today = datetime.today()
    date_string = today.strftime('%Y-%m-%d')

    base_name = 'format-compiled-doc-' if format else 'simple-compiled-doc-'
    next_file = get_next_filename(output_folder, base_name + date_string, 'docx')
    output_file = os.path.join(output_folder, next_file)
    doc.save(output_file)

    print("The Docx was saved!")

    # Open the saved document using the default application ('Pages' for Mac)
    # subprocess.run(['open', output_file])


def main():
    folder = '/Users/USER/Documents/github-clones/fraktur-ocr-transcription/output-txt/all-merged'
    output_folder = '/Users/USER/Documents/github-clones/fraktur-ocr-transcription/output-txt/all-merged/docx'
    os.makedirs(output_folder, exist_ok=True)

    process_choice = input("Create single docx or docx split in multiple parts? (single/parts): ")
    image_choice = input("Create with image insertion? (img folder needs to be in the script!) (y/n): ")

    format_choice = input("Make the Docx with formatting? (y/n/both): ")
    options = {
        "y": True,
        "n": False, 
        "single": False,
        "parts": True
    }

    if format_choice == "both":
        txt_to_word(folder, output_folder, format=False, parts=options[process_choice], img_available=options[image_choice])
        txt_to_word(folder, output_folder, format=True, parts=options[process_choice], img_available=options[image_choice])
    else:
        txt_to_word(folder, output_folder, options[format_choice], parts=options[process_choice], img_available=options[image_choice])


if __name__ == "__main__":
    main()
