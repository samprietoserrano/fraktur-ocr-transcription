import os
import re
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Inches
from PIL import Image, ImageOps

def add_hyperlink(doc, paragraph, text, url):
    """
    FOR DOCX FORMATTING.
    Adds a hyperlink to the given paragraph in the document.
    """
    part = doc.part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    r = paragraph.add_run(text)
    r.font.color.rgb = RGBColor(0, 0, 255)
    r.font.underline = True
    r.hyperlink = part


def apply_formatting(paragraph, text):
    """
    FOR TEXT FORMATTING.
    Adds color/bold/italics to words marked as 'still unknown' during spellchecking.
    """
    if "**" in text:
        parts = text.split("**")
        for i, part in enumerate(parts):
            if i % 2 == 0:
                paragraph.add_run(part)
            else:
                run = paragraph.add_run(part)
                run.font.color.rgb = RGBColor(0xdd, 0x2b, 0x05)
                run.font.underline = True
                run.bold = True
    else:
        paragraph.add_run(text)


WNS_COLS_NUM = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"

def set_number_of_columns(section, number):
    """
    FOR TEXT FORMATTING.
    Makes provided section have 2 columns using xpath.
    """
    section._sectPr.xpath("./w:cols")[0].set(WNS_COLS_NUM, str(number))


def modified_path(image_path):
    """
    FOR IMAGE PROCESSING.
    Make new path for the modified image. 
    """
    directory, filename = os.path.split(image_path)

    # Create the "images-only" folder if doesn't exists in the current directory (for organization)
    images_only_folder = os.path.join(directory, "images-only")
    if not os.path.exists(images_only_folder):
        os.makedirs(images_only_folder, exist_ok=True)

    new_folder = os.path.join(images_only_folder, "with_borders")
    os.makedirs(new_folder, exist_ok=True)
    return os.path.join(new_folder, filename) # return path of modified image


def add_border_to_image(input_path, border):
    """
    FOR IMAGE PROCESSING.
    Add color border to image, then save. Return path of modified image. 
    """
    # Modify the image
    img = Image.open(input_path)
    if isinstance(border, int) or isinstance(border, tuple):
        bordered_img = ImageOps.expand(img, border=border, fill='black')
    else:
        raise RuntimeError('Border is not an image or tuple')
    
    # Save new image, and return the path
    output_path = modified_path(input_path)
    bordered_img.save(output_path)
    return output_path


def add_image_to_page(document_obj, image_path, file_number, type):
    """
    FOR IMAGE PROCESSING.
    Insert modified image centered on path. 
    """
    if type == "img":
        suffix = '-IMAGE'
    else:
        suffix = '-TABLE'

    image_heading = str(file_number) + suffix
    image_section = document_obj.add_section(WD_SECTION.NEW_PAGE)
    set_number_of_columns(image_section, 1)

    document_obj.add_heading(image_heading, level=1)

    new_image_path = add_border_to_image(image_path, border=5)

    document_obj.add_picture(new_image_path, width=Inches(5), height=Inches(8))

    last_paragraph = document_obj.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def include_image_files(txt_files, img_nums):
    """
    FOR IMAGE PROCESSING.
    Merges the list of txt files and img files. 
    """
    img_and_txt = []
    for num in img_nums:
        img_and_txt.append(f'{num:02}.txt')
        # if img_and_txt[-1] in txt_files:
            # print("in both: " + img_and_txt[-1])

    img_and_txt += txt_files
    return sorted(img_and_txt)


def extract_number(filename):
    """Extract numbers from the filename using regular expressions."""
    match = re.search(r'\d+', filename)
    return int(match.group()) if match else float('inf')


def get_next_filename(folder, base_name, extension):
    """Generate the next available filename with an incrementing suffix."""
    counter = 1
    while True:
        filename = f"{base_name}{f'_{counter:02}' if counter > 1 else ''}.{extension}"
        filepath = os.path.join(folder, filename)
        if not os.path.exists(filepath):
            return filename
        counter += 1


def find_matching_path(folder, number):
    """Generate the image filepath for the matching filenumber."""
    for filename in os.listdir(folder):
        if filename.endswith(".jpg"):
            num_in_file = extract_number(filename)
            if num_in_file == number:
                path = os.path.join(folder, filename)
                return path

